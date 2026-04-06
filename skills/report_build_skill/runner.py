#!/usr/bin/env python3
import argparse
import json
import sys
from pathlib import Path
from datetime import datetime, date, timezone
from decimal import Decimal

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "scripts"))

from lib.insight_db import db_connect, set_search_path
from lib.insight_skill_runtime import build_standard_output, load_request

SKILL_NAME = "report_build_skill"

DEFAULT_REQUEST = {
    "analysis_type": "company_profile",
    "target_company": "Canonical",
    "domain": "canonical.com",
    "artifact_types": ["word", "excel", "ppt"]
}

def utc_now():
    return datetime.now(timezone.utc)

def ts_compact():
    return utc_now().strftime("%Y%m%d_%H%M%S")

def json_safe(value):
    if isinstance(value, dict):
        return {k: json_safe(v) for k, v in value.items()}
    if isinstance(value, list):
        return [json_safe(v) for v in value]
    if isinstance(value, tuple):
        return [json_safe(v) for v in value]
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, Decimal):
        return float(value)
    return value

def fetch_one(cur, sql, params=()):
    cur.execute(sql, params)
    row = cur.fetchone()
    if not row:
        return None
    cols = [d[0] for d in cur.description]
    return dict(zip(cols, row))

def fetch_all(cur, sql, params=()):
    cur.execute(sql, params)
    cols = [d[0] for d in cur.description]
    return [dict(zip(cols, row)) for row in cur.fetchall()]

def ensure_output_dir(company_name: str) -> Path:
    slug = "".join(c.lower() if c.isalnum() else "_" for c in company_name).strip("_") or "company"
    out_dir = ROOT / "outputs" / "report_build" / slug / ts_compact()
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir

def write_word_report(path: Path, company: dict, report_sections: dict, analysis_runs, snapshots, evidence_rows, metric_rows):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    doc.add_heading(f"{company.get('company_name')} — ORIS DB-backed Report", level=1)

    doc.add_heading("结论", level=2)
    doc.add_paragraph(report_sections["conclusion"])

    doc.add_heading("核心数据", level=2)
    for item in report_sections["core_data"]:
        doc.add_paragraph(f"{item['field']}: {item['value']}", style="List Bullet")

    doc.add_heading("数据出处", level=2)
    for item in report_sections["sources"]:
        line = f"{item.get('source_name')} | {item.get('source_type')} | {item.get('url')}"
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("分析与推断", level=2)
    for text in report_sections["analysis"]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("主要风险", level=2)
    for text in report_sections["risks"]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("待验证点 / 下一步", level=2)
    for text in report_sections["next_steps"]:
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("附录：DB counts", level=2)
    doc.add_paragraph(f"analysis_runs: {len(analysis_runs)}")
    doc.add_paragraph(f"source_snapshots: {len(snapshots)}")
    doc.add_paragraph(f"evidence_items: {len(evidence_rows)}")
    doc.add_paragraph(f"metric_observations: {len(metric_rows)}")

    doc.save(path)

def write_excel_report(path: Path, company: dict, analysis_runs, snapshots, evidence_rows, metric_rows):
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "summary"
    ws1.append(["field", "value"])
    ws1.append(["company_id", company.get("id")])
    ws1.append(["company_code", company.get("company_code")])
    ws1.append(["company_name", company.get("company_name")])
    ws1.append(["domain", company.get("domain")])
    ws1.append(["region", company.get("region")])
    ws1.append(["analysis_run_count", len(analysis_runs)])
    ws1.append(["source_snapshot_count", len(snapshots)])
    ws1.append(["evidence_item_count", len(evidence_rows)])
    ws1.append(["metric_observation_count", len(metric_rows)])

    ws2 = wb.create_sheet("analysis_runs")
    ws2.append(["id", "run_code", "request_id", "analysis_type", "target_company_id"])
    for r in analysis_runs:
        ws2.append([
            r.get("id"), r.get("run_code"), r.get("request_id"),
            r.get("analysis_type"), r.get("target_company_id")
        ])

    ws3 = wb.create_sheet("snapshots")
    ws3.append([
        "id", "source_id", "company_id", "snapshot_type", "snapshot_title",
        "snapshot_url", "raw_storage_path", "parsed_text_storage_path",
        "created_at", "source_name", "source_type", "root_domain", "official_flag"
    ])
    for r in snapshots:
        ws3.append([
            r.get("id"), r.get("source_id"), r.get("company_id"), r.get("snapshot_type"),
            r.get("snapshot_title"), r.get("snapshot_url"), r.get("raw_storage_path"),
            r.get("parsed_text_storage_path"), str(r.get("created_at")),
            r.get("source_name"), r.get("source_type"), r.get("root_domain"), r.get("official_flag")
        ])

    ws4 = wb.create_sheet("evidence")
    ws4.append([
        "id", "source_snapshot_id", "company_id", "evidence_type", "evidence_title",
        "evidence_text", "evidence_number", "evidence_unit", "evidence_date", "confidence_score"
    ])
    for r in evidence_rows:
        ws4.append([
            r.get("id"), r.get("source_snapshot_id"), r.get("company_id"), r.get("evidence_type"),
            r.get("evidence_title"), r.get("evidence_text"), r.get("evidence_number"),
            r.get("evidence_unit"), str(r.get("evidence_date")), r.get("confidence_score")
        ])

    ws5 = wb.create_sheet("metrics")
    ws5.append([
        "id", "company_id", "metric_code", "metric_name", "metric_value", "metric_unit",
        "period_type", "observation_date", "source_snapshot_id", "evidence_item_id"
    ])
    for r in metric_rows:
        ws5.append([
            r.get("id"), r.get("company_id"), r.get("metric_code"), r.get("metric_name"),
            r.get("metric_value"), r.get("metric_unit"), r.get("period_type"),
            str(r.get("observation_date")), r.get("source_snapshot_id"), r.get("evidence_item_id")
        ])

    wb.save(path)

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-file", default=None)
    ap.add_argument("--input-json", default=None)
    args = ap.parse_args()

    request = load_request(args.input_file, args.input_json, DEFAULT_REQUEST)
    target_company = request.get("target_company") or request.get("company_name") or request.get("entity") or "unknown"
    domain = request.get("domain")
    artifact_types = request.get("artifact_types") or ["word", "excel", "ppt"]
    analysis_type = request.get("analysis_type") or "company_profile"

    conn = db_connect()
    try:
        with conn:
            with conn.cursor() as cur:
                set_search_path(cur)

                company = None
                if domain:
                    company = fetch_one(
                        cur,
                        """
                        SELECT id, company_code, company_name, domain, region, status
                        FROM company
                        WHERE domain=%s
                        ORDER BY id DESC
                        LIMIT 1
                        """,
                        (domain,),
                    )

                if not company:
                    company = fetch_one(
                        cur,
                        """
                        SELECT id, company_code, company_name, domain, region, status
                        FROM company
                        WHERE company_name=%s
                        ORDER BY id DESC
                        LIMIT 1
                        """,
                        (target_company,),
                    )

                if not company:
                    out = build_standard_output(
                        skill_name=SKILL_NAME,
                        request=request,
                        conclusion="report_build_skill could not find target company in insight DB yet.",
                        core_data=[
                            {"field": "target_company", "value": target_company},
                            {"field": "domain", "value": domain},
                            {"field": "db_company_found", "value": False}
                        ],
                        sources=[],
                        facts=[
                            "No matching company row was found in insight.company for the requested target."
                        ],
                        inferences=[
                            "report_build_skill depends on upstream ingest/profile data; official_source_ingest_skill should run first."
                        ],
                        hypotheses=[],
                        risks=[
                            "Without DB-backed company/evidence/metric rows, report assembly will drift into request-only narrative."
                        ],
                        next_steps=[
                            "Run official_source_ingest_skill.",
                            "Run company_profile_skill.",
                            "Re-run report_build_skill."
                        ],
                        source_plan=[],
                        db_write_plan=[],
                        artifact_plan=[]
                    )
                    print(json.dumps(json_safe(out), ensure_ascii=False, indent=2))
                    return

                company_id = company["id"]

                analysis_runs = fetch_all(
                    cur,
                    """
                    SELECT id, run_code, request_id, analysis_type, target_company_id
                    FROM analysis_run
                    WHERE target_company_id=%s
                    ORDER BY id DESC
                    LIMIT 20
                    """,
                    (company_id,),
                )

                snapshots = fetch_all(
                    cur,
                    """
                    SELECT
                        ss.id,
                        ss.source_id,
                        ss.company_id,
                        ss.snapshot_type,
                        ss.snapshot_title,
                        ss.snapshot_url,
                        ss.raw_storage_path,
                        ss.parsed_text_storage_path,
                        ss.created_at,
                        s.source_name,
                        s.source_type,
                        s.root_domain,
                        s.official_flag
                    FROM source_snapshot ss
                    LEFT JOIN source s ON s.id = ss.source_id
                    WHERE ss.company_id=%s
                    ORDER BY ss.id DESC
                    LIMIT 20
                    """,
                    (company_id,),
                )

                evidence_rows = fetch_all(
                    cur,
                    """
                    SELECT
                        id,
                        source_snapshot_id,
                        company_id,
                        evidence_type,
                        evidence_title,
                        evidence_text,
                        evidence_number,
                        evidence_unit,
                        evidence_date,
                        confidence_score
                    FROM evidence_item
                    WHERE company_id=%s
                    ORDER BY id DESC
                    LIMIT 20
                    """,
                    (company_id,),
                )

                metric_rows = fetch_all(
                    cur,
                    """
                    SELECT
                        id,
                        company_id,
                        metric_code,
                        metric_name,
                        metric_value,
                        metric_unit,
                        period_type,
                        observation_date,
                        source_snapshot_id,
                        evidence_item_id
                    FROM metric_observation
                    WHERE company_id=%s
                    ORDER BY id DESC
                    LIMIT 20
                    """,
                    (company_id,),
                )

                latest_run = analysis_runs[0] if analysis_runs else None

                report_sections = {
                    "conclusion": f"{target_company} 已具备 DB-backed 报告组装基础，但当前证据仍以 bootstrap 级来源采集证据为主。",
                    "core_data": [
                        {"field": "company_id", "value": company_id},
                        {"field": "company_name", "value": company.get("company_name")},
                        {"field": "analysis_run_count", "value": len(analysis_runs)},
                        {"field": "source_snapshot_count", "value": len(snapshots)},
                        {"field": "evidence_item_count", "value": len(evidence_rows)},
                        {"field": "metric_observation_count", "value": len(metric_rows)},
                    ],
                    "sources": [
                        {
                            "source_snapshot_id": s.get("id"),
                            "source_name": s.get("source_name"),
                            "source_type": s.get("source_type") or s.get("snapshot_type"),
                            "url": s.get("snapshot_url"),
                            "official_flag": s.get("official_flag"),
                        }
                        for s in snapshots
                    ],
                    "analysis": [
                        "当前 report_build_skill 已可直接消费 insight DB 中的 company / analysis_run / source_snapshot / evidence_item / metric_observation。",
                        "这意味着后续 Word / Excel / PPT 生成可以建立在持久化数据之上，而不是临时请求结果之上。"
                    ],
                    "risks": [
                        "当前 evidence_item 与 metric_observation 仍是 bootstrap-level placeholder。",
                        "尚未引入 citation_link，正式报告的逐条引用绑定还未闭环。",
                        "尚未做 snapshot 去重 / latest-first condensation，后续报告需控制重复来源噪音。"
                    ],
                    "next_steps": [
                        "给 official_source_ingest_skill 增加真实网页正文抓取。",
                        "从 parsed text 中抽取正文级 evidence_item。",
                        "补 citation_link。",
                        "再把 report_build_skill 接到 Word / Excel 物料生成。"
                    ]
                }

                out_dir = ensure_output_dir(company.get("company_name") or target_company)
                word_path = out_dir / f"{company.get('company_name','company').lower()}_db_report.docx"
                excel_path = out_dir / f"{company.get('company_name','company').lower()}_db_report.xlsx"
                json_path = out_dir / f"{company.get('company_name','company').lower()}_db_report.json"

                generated_artifacts = []

                if "word" in artifact_types:
                    write_word_report(word_path, company, report_sections, analysis_runs, snapshots, evidence_rows, metric_rows)
                    generated_artifacts.append({
                        "artifact_type": "word",
                        "path": str(word_path.relative_to(ROOT))
                    })

                if "excel" in artifact_types:
                    write_excel_report(excel_path, company, analysis_runs, snapshots, evidence_rows, metric_rows)
                    generated_artifacts.append({
                        "artifact_type": "excel",
                        "path": str(excel_path.relative_to(ROOT))
                    })

                report_json = {
                    "company": company,
                    "latest_analysis_run": latest_run,
                    "analysis_runs": analysis_runs,
                    "recent_snapshots": snapshots,
                    "recent_evidence_items": evidence_rows,
                    "recent_metric_observations": metric_rows,
                    "report_sections": report_sections,
                    "generated_artifacts": generated_artifacts,
                }

                json_path.write_text(json.dumps(json_safe(report_json), ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
                generated_artifacts.append({
                    "artifact_type": "json",
                    "path": str(json_path.relative_to(ROOT))
                })

                facts = [
                    f"Target company found in insight.company with company_id={company_id}.",
                    f"Recent analysis_run count in report view: {len(analysis_runs)}.",
                    f"Recent source_snapshot count in report view: {len(snapshots)}.",
                    f"Recent evidence_item count in report view: {len(evidence_rows)}.",
                    f"Recent metric_observation count in report view: {len(metric_rows)}.",
                    f"Generated local artifact count: {len(generated_artifacts)}."
                ]

                inferences = [
                    "The DB-backed report input layer now exists.",
                    "Word/Excel local artifact generation is now attached to persisted report input."
                ]

                out = build_standard_output(
                    skill_name=SKILL_NAME,
                    request=request,
                    conclusion="report_build_skill now generates local DB-backed Word/Excel artifacts from persisted insight data.",
                    core_data=[
                        {"field": "company_id", "value": company_id},
                        {"field": "company_name", "value": company.get("company_name")},
                        {"field": "analysis_type", "value": analysis_type},
                        {"field": "artifact_types_requested", "value": artifact_types},
                        {"field": "latest_analysis_run_id", "value": latest_run.get("id") if latest_run else None},
                        {"field": "source_snapshot_count_recent", "value": len(snapshots)},
                        {"field": "evidence_item_count_recent", "value": len(evidence_rows)},
                        {"field": "metric_observation_count_recent", "value": len(metric_rows)},
                        {"field": "generated_artifact_count", "value": len(generated_artifacts)},
                    ],
                    sources=report_sections["sources"],
                    facts=facts,
                    inferences=inferences,
                    hypotheses=[
                        "Once citation_link and body-level evidence are added, these local artifacts can be promoted into formal delivery artifacts."
                    ],
                    risks=report_sections["risks"],
                    next_steps=[
                        "Verify generated Word/Excel files open correctly.",
                        "Register artifacts into insight.report_artifact.",
                        "Create delivery_task rows for Feishu/Qbot distribution."
                    ],
                    source_plan=snapshots,
                    db_write_plan=[],
                    artifact_plan=generated_artifacts,
                )

                out["db_backed_report"] = report_json

                print(json.dumps(json_safe(out), ensure_ascii=False, indent=2))
    finally:
        conn.close()

if __name__ == "__main__":
    main()
