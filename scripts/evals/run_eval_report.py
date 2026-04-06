#!/usr/bin/env python3
import argparse
import json
import os
import re
import sys
import uuid
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import requests
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[2]
CFG_PATH = ROOT / "config" / "report_runtime.json"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def cfg_get(data, keys, default=None):
    cur = data
    for key in keys:
        if not isinstance(cur, dict) or key not in cur:
            return default
        cur = cur[key]
    return cur


def expand_path(root: Path, value: str) -> Path:
    p = Path(os.path.expanduser(value))
    if p.is_absolute():
        return p
    return root / value


def now_utc():
    return datetime.now(timezone.utc)


def ts_compact():
    return now_utc().strftime("%Y%m%d_%H%M%S")


def normalize_text(text):
    return re.sub(r"\s+", " ", str(text or "")).strip()


def extract_urls(text):
    return re.findall(r"https?://[^\s\)]+", text or "", re.IGNORECASE)


def contains_numberish(text):
    patterns = [
        r"\d+(?:\.\d+)?%",
        r"\d{4}年",
        r"\d+(?:\.\d+)?亿",
        r"\d+(?:\.\d+)?万",
        r"\d+(?:\.\d+)?元",
        r"\d+(?:\.\d+)?倍",
        r"\d+(?:\.\d+)?月",
        r"\d+(?:\.\d+)?日",
        r"\d+"
    ]
    return any(re.search(p, text or "") for p in patterns)


def has_source_cue(text):
    cues = [
        "数据来源", "来源", "出处", "根据", "公告", "财报", "年报", "季报",
        "国家统计局", "交易所", "公司公告", "source", "sources",
        "according to", "filing", "report"
    ]
    low = (text or "").lower()
    return any(c.lower() in low for c in cues)


def has_fact_infer_risk_structure(text):
    cues = ["事实", "推断", "假设", "风险", "结论", "判断", "验证", "不确定性"]
    return sum(1 for c in cues if c in (text or "")) >= 2


def has_critical_reasoning(text):
    cues = [
        "但是", "不过", "与此同时", "另一方面", "前提是", "取决于", "边际",
        "增量", "验证", "证伪", "条件", "若", "则", "因为", "所以",
        "风险在于", "核心变量", "约束", "敏感性", "情景"
    ]
    return sum(1 for c in cues if c in (text or "")) >= 3


def mask_secret(value):
    if not value:
        return ""
    s = str(value)
    if len(s) <= 8:
        return "*" * len(s)
    return s[:6] + "*" * (len(s) - 10) + s[-4:]


def normalize_key_name(name: str) -> str:
    return re.sub(r"[^a-z0-9]", "", str(name).lower())


def recursive_find_secret(obj, normalized_targets):
    if isinstance(obj, dict):
        for k, v in obj.items():
            nk = normalize_key_name(k)
            if nk in normalized_targets and isinstance(v, str) and v.strip():
                return v.strip(), k
        for _, v in obj.items():
            found = recursive_find_secret(v, normalized_targets)
            if found:
                return found
    elif isinstance(obj, list):
        for item in obj:
            found = recursive_find_secret(item, normalized_targets)
            if found:
                return found
    return None


def resolve_api_key(cfg, cli_api_key=None):
    auth_cfg = cfg_get(cfg, ["oris_api", "auth"], {}) or {}
    env_var = auth_cfg.get("env_var")
    secrets_json_path = expand_path(ROOT, auth_cfg.get("secrets_json_path", "~/.openclaw/secrets.json"))
    discovery_keys = auth_cfg.get("discovery_keys", [])

    if cli_api_key:
        return {
            "ok": True,
            "api_key": cli_api_key,
            "source": "cli_argument"
        }

    if env_var and os.environ.get(env_var):
        return {
            "ok": True,
            "api_key": os.environ.get(env_var),
            "source": f"env:{env_var}"
        }

    if secrets_json_path.exists():
        try:
            data = load_json(secrets_json_path)
            normalized_targets = {normalize_key_name(x) for x in discovery_keys}
            found = recursive_find_secret(data, normalized_targets)
            if found:
                value, hit_key = found
                return {
                    "ok": True,
                    "api_key": value,
                    "source": f"secrets_json:{secrets_json_path}::{hit_key}"
                }
        except Exception as e:
            return {
                "ok": False,
                "error": f"failed_to_parse_secrets_json: {e}",
                "source": str(secrets_json_path)
            }

    return {
        "ok": False,
        "error": "api_key_not_found",
        "checked_env_var": env_var,
        "checked_secrets_json": str(secrets_json_path),
        "discovery_keys": discovery_keys
    }


def call_oris(cfg, api_key, question, role):
    infer_url = cfg_get(cfg, ["oris_api", "infer_url"])
    timeout_seconds = int(cfg_get(cfg, ["oris_api", "timeout_seconds"], 300))
    header_name = cfg_get(cfg, ["oris_api", "auth", "header_name"], "X-ORIS-API-Key")
    default_source = cfg_get(cfg, ["report", "default_source"], "eval_report_batch")

    headers = {
        "Content-Type": "application/json",
        header_name: api_key
    }
    payload = {
        "role": role,
        "prompt": question,
        "source": default_source,
        "request_id": str(uuid.uuid4())
    }

    resp = requests.post(infer_url, headers=headers, json=payload, timeout=timeout_seconds)
    resp.raise_for_status()
    data = resp.json()

    if not data.get("ok"):
        return {
            "ok": False,
            "error": json.dumps(data, ensure_ascii=False)
        }

    d = data.get("data") or {}
    return {
        "ok": True,
        "request_id": data.get("request_id"),
        "role": d.get("role"),
        "selected_model": d.get("selected_model"),
        "execution_primary": d.get("execution_primary"),
        "used_provider": d.get("used_provider"),
        "used_model": d.get("used_model"),
        "attempt": d.get("attempt"),
        "text": d.get("text", "")
    }


def score_answer(answer_text, question_text):
    text = answer_text or ""
    q = question_text or ""

    relevance = 5 if any(token in text for token in re.findall(r"[\u4e00-\u9fffA-Za-z0-9]{2,}", q)[:10]) else 3
    if len(normalize_text(text)) < 20:
        relevance = 1

    logic = 5 if has_critical_reasoning(text) else 3
    if any(x in text for x in ["因为", "所以", "结论", "判断", "原因"]):
        logic = max(logic, 4)

    data_support = 5 if (contains_numberish(text) and has_source_cue(text)) else 3 if contains_numberish(text) else 1
    source_transparency = 5 if (has_source_cue(text) and extract_urls(text)) else 3 if has_source_cue(text) else 1
    decision_utility = 5 if has_fact_infer_risk_structure(text) else 3
    safety_restraint = 5 if any(x in text for x in ["可能", "大概率", "需要继续验证", "不确定", "取决于"]) else 3

    dim = {
        "relevance": relevance,
        "logic": logic,
        "data_support": data_support,
        "source_transparency": source_transparency,
        "decision_utility": decision_utility,
        "safety_restraint": safety_restraint
    }
    total_score = sum(dim.values())
    avg_score = round(total_score / len(dim), 2)

    return {
        "dimensions": dim,
        "total_score": total_score,
        "avg_score": avg_score,
        "evidence_flags": {
            "has_numeric_data": contains_numberish(text),
            "has_source_cue": has_source_cue(text),
            "has_explicit_link": bool(extract_urls(text)),
            "has_fact_infer_risk": has_fact_infer_risk_structure(text),
            "source_links": "\n".join(extract_urls(text)[:20])
        }
    }


def build_summary(results):
    if not results:
        return {
            "count": 0,
            "avg_total_score": 0,
            "avg_score": 0,
            "data_ratio": 0,
            "source_ratio": 0,
            "link_ratio": 0,
            "fact_risk_ratio": 0
        }

    count = len(results)
    avg_total_score = round(sum(r["score"]["total_score"] for r in results) / count, 2)
    avg_score = round(sum(r["score"]["avg_score"] for r in results) / count, 2)

    def ratio(flag_name):
        n = sum(1 for r in results if r["score"]["evidence_flags"][flag_name])
        return round(n * 100 / count, 2)

    return {
        "count": count,
        "avg_total_score": avg_total_score,
        "avg_score": avg_score,
        "data_ratio": ratio("has_numeric_data"),
        "source_ratio": ratio("has_source_cue"),
        "link_ratio": ratio("has_explicit_link"),
        "fact_risk_ratio": ratio("has_fact_infer_risk")
    }


def write_excel(results, out_path: Path):
    wb = Workbook()

    ws = wb.active
    ws.title = "scoring"
    ws.append([
        "case_id", "category", "role", "selected_model", "used_provider", "used_model",
        "relevance", "logic", "data_support", "source_transparency",
        "decision_utility", "safety_restraint", "total_score", "avg_score"
    ])

    for r in results:
        d = r["score"]["dimensions"]
        ws.append([
            r["case_id"], r["category"], r["role"], r.get("selected_model"),
            r.get("used_provider"), r.get("used_model"),
            d["relevance"], d["logic"], d["data_support"], d["source_transparency"],
            d["decision_utility"], d["safety_restraint"], r["score"]["total_score"], r["score"]["avg_score"]
        ])

    ws2 = wb.create_sheet("raw_outputs")
    ws2.append([
        "case_id", "question", "answer_text", "request_id", "source_links"
    ])
    for r in results:
        ws2.append([
            r["case_id"], r["question"], r["answer_text"], r.get("request_id"),
            r["score"]["evidence_flags"]["source_links"]
        ])

    ws3 = wb.create_sheet("evidence_check")
    ws3.append([
        "case_id", "has_numeric_data", "has_source_cue",
        "has_explicit_link", "has_fact_infer_risk"
    ])
    for r in results:
        e = r["score"]["evidence_flags"]
        ws3.append([
            r["case_id"], str(e["has_numeric_data"]), str(e["has_source_cue"]),
            str(e["has_explicit_link"]), str(e["has_fact_infer_risk"])
        ])

    fill = PatternFill(fill_type="solid", fgColor="D9EAF7")
    bold = Font(bold=True)

    for sheet in [ws, ws2, ws3]:
        for cell in sheet[1]:
            cell.font = bold
            cell.fill = fill
            cell.alignment = Alignment(vertical="top", wrap_text=True)
        sheet.freeze_panes = "A2"
        for col in sheet.columns:
            max_len = max(len(str(c.value)) if c.value is not None else 0 for c in col)
            sheet.column_dimensions[col[0].column_letter].width = min(max(max_len + 2, 12), 60)

    wb.save(out_path)


def write_docx(cfg, results, summary, out_path: Path):
    title = cfg_get(cfg, ["report", "title"], "ORIS 评测主报告")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"生成时间：{now_utc().astimezone().strftime('%Y-%m-%d %H:%M:%S %Z')}")
    doc.add_paragraph(f"总题数：{summary['count']}")
    doc.add_paragraph(f"平均总分（满分30）：{summary['avg_total_score']}")
    doc.add_paragraph(f"平均单项分：{summary['avg_score']}")

    doc.add_heading("一、执行摘要", level=1)
    doc.add_paragraph(f"有明确数据支撑的题目占比：{summary['data_ratio']}%")
    doc.add_paragraph(f"有明确来源提示的题目占比：{summary['source_ratio']}%")
    doc.add_paragraph(f"有明确链接的题目占比：{summary['link_ratio']}%")
    doc.add_paragraph(f"有事实/推断/风险结构的题目占比：{summary['fact_risk_ratio']}%")

    doc.add_heading("二、判定规则", level=1)
    doc.add_paragraph("本轮评测重点检查：回答是否说明了数据是什么、数据出处是什么、是否给出明确链接，以及是否将事实、推断、风险分层表达。")

    doc.add_heading("三、逐题结果", level=1)
    for r in results:
        doc.add_heading(f"{r['case_id']}｜{r['category']}", level=2)
        doc.add_paragraph(f"问题：{r['question']}")
        doc.add_paragraph(f"角色：{r['role']}｜模型：{r.get('used_model') or ''}｜Provider：{r.get('used_provider') or ''}")
        doc.add_paragraph(
            f"评分：相关性 {r['score']['dimensions']['relevance']}，"
            f"逻辑性 {r['score']['dimensions']['logic']}，"
            f"数据支撑 {r['score']['dimensions']['data_support']}，"
            f"来源透明 {r['score']['dimensions']['source_transparency']}，"
            f"决策可用 {r['score']['dimensions']['decision_utility']}，"
            f"审慎性 {r['score']['dimensions']['safety_restraint']}，"
            f"总分 {r['score']['total_score']}/30"
        )
        doc.add_paragraph(f"回答：{r['answer_text']}")
        flags = r["score"]["evidence_flags"]
        doc.add_paragraph(
            f"检查：数据={flags['has_numeric_data']}；来源提示={flags['has_source_cue']}；"
            f"明确链接={flags['has_explicit_link']}；事实/推断/风险结构={flags['has_fact_infer_risk']}"
        )
        if flags["source_links"]:
            doc.add_paragraph(f"抓取到的链接/出处：{flags['source_links']}")

    doc.add_heading("四、改进建议", level=1)
    doc.add_paragraph("1. 对需要判断和结论的回答，默认补充：数据是什么、口径是什么、时间点是什么、出处是什么、链接是什么。")
    doc.add_paragraph("2. 对涉及预测、趋势、估值、政策、竞争格局的回答，默认按：事实 / 推断 / 风险 / 待验证点 输出。")
    doc.add_paragraph("3. 报告类问题优先保证正确性和证据链完整，再考虑响应速度。")

    doc.save(out_path)


def write_manifest(cfg, generated_files, out_path: Path):
    manifest = {
        "version": 1,
        "generated_at": now_utc().isoformat(),
        "download_targets": cfg_get(cfg, ["distribution", "download_targets"], []),
        "files": []
    }
    for file_path in generated_files:
        manifest["files"].append({
            "filename": file_path.name,
            "absolute_path": str(file_path.resolve()),
            "size_bytes": file_path.stat().st_size if file_path.exists() else 0
        })
    out_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")


def write_zip(bundle_path: Path, files):
    with zipfile.ZipFile(bundle_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in files:
            if p.exists():
                zf.write(p, arcname=p.name)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--api-key", default=None)
    parser.add_argument("--check-auth", action="store_true")
    args = parser.parse_args()

    cfg = load_json(CFG_PATH)
    out_dir = expand_path(ROOT, cfg_get(cfg, ["paths", "output_dir"]))
    bank_path = expand_path(ROOT, cfg_get(cfg, ["paths", "eval_bank"]))
    out_dir.mkdir(parents=True, exist_ok=True)

    auth = resolve_api_key(cfg, cli_api_key=args.api_key)

    if args.check_auth:
        print(json.dumps({
            "ok": auth.get("ok", False),
            "auth_source": auth.get("source"),
            "api_key_masked": mask_secret(auth.get("api_key")) if auth.get("api_key") else None,
            "error": auth.get("error"),
            "checked_env_var": auth.get("checked_env_var"),
            "checked_secrets_json": auth.get("checked_secrets_json")
        }, ensure_ascii=False, indent=2))
        return

    if not auth.get("ok"):
        print(json.dumps(auth, ensure_ascii=False, indent=2))
        sys.exit(1)

    if not bank_path.exists():
        print(f"missing eval bank: {bank_path}")
        sys.exit(1)

    bank = load_json(bank_path)
    cases = bank.get("cases") if isinstance(bank, dict) else bank
    if not cases:
        print("eval bank is empty")
        sys.exit(1)

    prefix = cfg_get(cfg, ["report", "file_prefix"], "oris_eval")
    ts = ts_compact()

    xlsx_path = out_dir / f"{prefix}_scoring_{ts}.xlsx"
    docx_path = out_dir / f"{prefix}_report_{ts}.docx"
    manifest_path = out_dir / f"{prefix}_download_manifest_{ts}.json"
    bundle_prefix = cfg_get(cfg, ["distribution", "bundle_prefix"], "oris_eval_package")
    zip_path = out_dir / f"{bundle_prefix}_{ts}.zip"

    default_role = cfg_get(cfg, ["report", "default_role"], "primary_general")

    results = []
    for idx, case in enumerate(cases, start=1):
        case_id = case.get("case_id") or case.get("id") or f"CASE_{idx:03d}"
        category = case.get("category") or "general"
        role = case.get("role") or default_role
        question = case.get("question") or case.get("prompt") or ""

        if not normalize_text(question):
            continue

        print(f"[{idx}/{len(cases)}] {case_id}")

        try:
            infer_result = call_oris(cfg, auth["api_key"], question, role)
        except Exception as e:
            infer_result = {
                "ok": False,
                "error": str(e)
            }

        answer_text = infer_result.get("text", "") if infer_result.get("ok") else f"[ERROR] {infer_result.get('error')}"
        score = score_answer(answer_text, question)

        results.append({
            "case_id": case_id,
            "category": category,
            "role": role,
            "question": question,
            "request_id": infer_result.get("request_id"),
            "selected_model": infer_result.get("selected_model"),
            "execution_primary": infer_result.get("execution_primary"),
            "used_provider": infer_result.get("used_provider"),
            "used_model": infer_result.get("used_model"),
            "attempt": infer_result.get("attempt"),
            "answer_text": answer_text,
            "score": score
        })

    summary = build_summary(results)

    generated_files = []

    if cfg_get(cfg, ["report", "emit", "xlsx"], True):
        write_excel(results, xlsx_path)
        generated_files.append(xlsx_path)

    if cfg_get(cfg, ["report", "emit", "docx"], True):
        write_docx(cfg, results, summary, docx_path)
        generated_files.append(docx_path)

    if cfg_get(cfg, ["report", "emit", "manifest"], True):
        write_manifest(cfg, generated_files, manifest_path)
        generated_files.append(manifest_path)

    if cfg_get(cfg, ["report", "emit", "zip_bundle"], True):
        write_zip(zip_path, generated_files)
        generated_files.append(zip_path)

    print()
    print("===== done =====")
    print(json.dumps({
        "auth_source": auth.get("source"),
        "generated_files": [str(p) for p in generated_files],
        "summary": summary
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
